"""
Dokumentasi Teknis - Catatan Keuangan
Professional DOCX with green theme styling
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import datetime
import os

# ── Color Theme (Forest Green) ───────────────────────────
C_DARK     = RGBColor(0x1B, 0x5E, 0x20)   # Deep forest - headings
C_PRIMARY  = RGBColor(0x2E, 0x7D, 0x32)   # Green - accent
C_SECOND   = RGBColor(0x58, 0x81, 0x57)   # Medium green
C_LIGHT    = RGBColor(0xA3, 0xB1, 0x8A)   # Sage - subtle
C_BG       = RGBColor(0xF8, 0xFC, 0xF8)   # Very light green bg
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY     = RGBColor(0x55, 0x55, 0x55)   # Body text
C_LGRAY    = RGBColor(0x88, 0x88, 0x88)   # Light gray
C_TBL_HDR  = RGBColor(0x2E, 0x7D, 0x32)
C_TBL_ALT  = RGBColor(0xF0, 0xF7, 0xF0)   # Alternating row

HEX_DARK   = "1B5E20"
HEX_PRIMARY = "2E7D32"
HEX_SECOND  = "588157"
HEX_LIGHT   = "A3B18A"
HEX_BG      = "F8FCF8"
HEX_WHITE   = "FFFFFF"
HEX_GRAY    = "555555"
HEX_LGRAY   = "888888"
HEX_ALT     = "F0F7F0"

# ── Helpers ──────────────────────────────────────────────
doc = None

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def para_spacing(para, before=0, after=6):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before * 20))
    spacing.set(qn('w:after'), str(after * 20))
    pPr.append(spacing)

def set_run_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

# ── Document Setup ────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = C_GRAY
pPr = style._element.get_or_add_pPr()
spacing = OxmlElement('w:spacing')
spacing.set(qn('w:after'), '80')
pPr.append(spacing)

# ── COVER PAGE ──────────────────────────────────────────────
# Top accent bar
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run(' ')
run.font.size = Pt(6)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '24')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), HEX_PRIMARY)
pBdr.append(bottom)
pPr.append(pBdr)

# Spacer
for _ in range(2):
    doc.add_paragraph()

# App name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p, 0, 4)
run = p.add_run('CATATAN KEUANGAN')
run.font.name = 'Calibri'
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = C_PRIMARY

# Tagline
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p, 0, 24)
run = p.add_run('Aplikasi Pencatatan Keuangan Personal')
run.font.name = 'Calibri'
run.font.size = Pt(20)
run.font.color.rgb = C_SECOND

# Divider line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p, 0, 24)
run = p.add_run('─' * 50)
run.font.color.rgb = C_LIGHT

# Info block
info = [
    ('Versi', '1.0.0'),
    ('Platform', 'Flutter / Android / iOS'),
    ('Tanggal', datetime.date.today().strftime('%d %B %Y')),
    ('Bahasa', 'Dart 3.11 / Flutter 3.41'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, 0, 4)
    run = p.add_run(f'{label}: ')
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = C_LGRAY
    run2 = p.add_run(value)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(12)
    run2.font.color.rgb = C_GRAY

# Bottom bar
for _ in range(1):
    doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run(' ')
run.font.size = Pt(4)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '24')
top.set(qn('w:space'), '1')
top.set(qn('w:color'), HEX_PRIMARY)
pBdr.append(top)
pPr.append(pBdr)

doc.add_page_break()

# ── Helper: Section Heading ────────────────────────────────
def h1(text):
    # Green bar + H1 text
    p = doc.add_paragraph()
    para_spacing(p, 20, 8)
    # Left accent bar via border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '36')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), HEX_PRIMARY)
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = C_DARK
    return p

def h2(text):
    p = doc.add_paragraph()
    para_spacing(p, 16, 4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = C_PRIMARY
    return p

def h3(text):
    p = doc.add_paragraph()
    para_spacing(p, 10, 2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = C_SECOND
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    para_spacing(p, 2, 6)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = C_GRAY
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    para_spacing(p, 1, 3)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = C_GRAY
    return p

def code(text):
    p = doc.add_paragraph()
    para_spacing(p, 4, 4)
    p.paragraph_format.left_indent = Inches(0.3)
    # Light gray bg
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4F0')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = C_DARK
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, HEX_PRIMARY)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = C_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            set_cell_bg(cell, HEX_ALT if ri % 2 == 0 else HEX_WHITE)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = C_GRAY

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return t

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════
h1('Daftar Isi')

toc_items = [
    ('1', 'Gambaran Umum', '3'),
    ('2', 'Tech Stack & Struktur', '4'),
    ('3', 'Database Schema', '4'),
    ('4', 'Fitur Utama', '5'),
    ('5', 'Services & Arsitektur', '7'),
    ('6', 'Keamanan & Enkripsi', '8'),
    ('7', 'Deep Linking & Widget', '9'),
    ('8', 'Testing & Build', '10'),
]

for num, title, pg in toc_items:
    p = doc.add_paragraph()
    para_spacing(p, 3, 3)
    run = p.add_run(f'{num}.  ')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = C_PRIMARY
    run2 = p.add_run(title)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(11)
    run2.font.color.rgb = C_GRAY
    run3 = p.add_run(f'  ...............  {pg}')
    run3.font.name = 'Calibri'
    run3.font.size = Pt(10)
    run3.font.color.rgb = C_LIGHT

page_break()

# ═══════════════════════════════════════════════════════════
# 2. GAMBARAN UMUM
# ═══════════════════════════════════════════════════════════
h1('1. Gambaran Umum')

body('Catatan Keuangan adalah aplikasi pencatatan keuangan personal yang dikembangkan dengan Flutter. '
     'Dirancang untuk membantu pengguna mengelola keuangan pribadi dengan fitur lengkap meliputi pencatatan '
     'transaksi, multi-dompet, budget bulanan, recurring transactions, laporan PDF dengan lampiran, '
     'backup terenkripsi AES-256, biometric authentication, dan Android home screen widget.')

h2('1.1 Informasi Project')
table(
    ['Properti', 'Nilai'],
    [
        ['Nama Project', 'catatan_keuangan'],
        ['Package', 'com.example.catatan_keuangan'],
        ['Versi', '1.0.0'],
        ['Flutter SDK', '3.41.x'],
        ['Dart SDK', '3.11.x'],
        ['Target Platform', 'Android (minSdk 21), iOS'],
        ['State Management', 'Riverpod 3.x'],
        ['Routing', 'GoRouter 14.x'],
    ],
    col_widths=[2.0, 4.5]
)

h2('1.2 Statistik Kode')
table(
    ['Metrik', 'Jumlah'],
    [
        ['Total baris kode (lib/)', '~2,800 baris'],
        ['Total file Dart', '23 file'],
        ['Test cases', '63 tests (3 test files)'],
        ['Dependencies', '15 package'],
        ['Database version', 'v8 (dengan migration)'],
        ['APK variants', '3 (armeabi, arm64, universal)'],
    ],
    col_widths=[3.0, 3.5]
)

# ═══════════════════════════════════════════════════════════
# 3. TECH STACK & STRUKTUR
# ═══════════════════════════════════════════════════════════
h1('2. Tech Stack & Struktur')

h2('2.1 Dependencies')
table(
    ['Kategori', 'Package', 'Versi', 'Fungsi'],
    [
        ['Database', 'sqflite', '^2.3.0', 'ORM SQLite'],
        ['State', 'flutter_riverpod', '^3.3.1', 'Manajemen state'],
        ['Navigation', 'go_router', '^14.8.1', 'Routing + deep linking'],
        ['PDF', 'pdf + printing', '3.10.8 / 5.12.0', 'Generate & share PDF'],
        ['File', 'image_picker + file_picker', 'various', 'Kamera, galeri, file picker'],
        ['Enkripsi', 'encrypt + pointycastle', '5.0.3 / 3.9.1', 'Enkripsi AES-256-CBC'],
        ['Biometric', 'local_auth', '^2.3.0', 'Fingerprint / Face ID'],
        ['Widget', 'home_widget', '^0.6.0', 'Widget layar utama Android'],
        ['Format', 'intl', '^0.20.2', 'Locale Indonesia'],
    ],
    col_widths=[1.2, 1.8, 1.3, 2.2]
)

h2('2.2 Struktur Direktori')
code('''lib/
├── main.dart                  Titik awal aplikasi
├── app.dart                   MaterialApp + overlay kunci
├── router.dart                 GoRouter + deep linking
├── providers.dart              Provider Riverpod
├── models/
│   ├── models.dart            Transaksi, Dompet, Kategori, Budget, Pengaturan
│   └── constants.dart         Konstanta aplikasi
├── services/
│   ├── database_helper.dart   CRUD SQLite + migrasi (v8)
│   ├── crypto_service.dart     AES-256-CBC + PBKDF2
│   ├── biometric_service.dart  Wrapper local_auth
│   ├── file_service.dart       Ambil gambar/file
│   ├── home_widget_service.dart Updater widget Android
│   └── pdf_laporan_service.dart Generator PDF
└── pages/
    ├── home_page.dart          TabBar (Hari Ini / Per Tanggal / Bulanan / Lainnya)
    ├── settings_page.dart      Mode gelap + toggle biometric
    ├── statistik_page.dart      Statistik pengeluaran
    ├── backup_page.dart        UI Backup & Restore
    ├── pin_lock_screen.dart    PIN + autentikasi biometric
    ├── sheets/                 Komponen bottom sheet
    └── widgets/                 Komponen widget reusable''')

# ═══════════════════════════════════════════════════════════
# 4. DATABASE SCHEMA
# ═══════════════════════════════════════════════════════════
h1('3. Database Schema')

body('SQLite dengan versi v8 dan sistem auto-migration. 6 index untuk query performa.')

h2('3.1 Tabel Database')
table(
    ['Tabel', 'Kolom Utama', 'Deskripsi'],
    [
        ['transaksi', 'id, jenis, jumlah, deskripsi, kategori,\ntanggal, lampiran, is_recurring,\nrecurring_frequency, id_dompet, deleted_at',
         'Transaksi dengan soft delete'],
        ['dompet', 'id, nama, saldo, warna', 'Multi-wallet dengan saldo'],
        ['kategori', 'id, nama, jenis, icon, is_default', 'Kategori dengan icon'],
        ['budget', 'id, bulan, tahun, nominal, kategori', 'Budget bulanan per kategori'],
        ['pengaturan', 'id, is_dark_mode, pin, use_biometric', 'Pengaturan aplikasi'],
    ],
    col_widths=[1.2, 3.3, 2.0]
)

h2('3.2 Index Database')
for idx in ['idx_transaksi_tanggal', 'idx_transaksi_kategori', 'idx_transaksi_jenis',
           'idx_transaksi_dompet', 'idx_budget_period', 'idx_budget_unique (UNIQUE)']:
    bullet(f'{idx}')

h2('3.3 Migration History')
table(
    ['Versi', 'Perubahan'],
    [
        ['v1–v3', 'Ditambah dompet, budget, pengaturan tables'],
        ['v4', 'Ditambah kolom lampiran ke transaksi'],
        ['v5', 'Ditambah kategori table'],
        ['v6', 'Ditambah index untuk performa query'],
        ['v7', 'Ditambah soft delete (deleted_at column)'],
        ['v8', 'Ditambah use_biometric ke pengaturan'],
    ],
    col_widths=[1.0, 5.5]
)

# ═══════════════════════════════════════════════════════════
# 5. FITUR UTAMA
# ═══════════════════════════════════════════════════════════
h1('4. Fitur Utama')

h2('4.1 Pencatatan Transaksi')
body('Fitur inti aplikasi untuk mencatat transaksi.')
bullet('Jenis: Pemasukan atau Pengeluaran')
bullet('Kategori: 9 default untuk pengeluaran + 6 untuk pemasukan, custom dengan icon')
bullet('Lampiran: Gambar dari kamera/galeri atau file arbitrary')
bullet('Recurring: Daily, Weekly, Monthly, Yearly — diproses otomatis saat app start')
bullet('Soft delete: Swipe kiri untuk hapus, dipindah ke tong sampah, bisa direstore')
bullet('Tap untuk edit, dismiss untuk hapus')

h2('4.2 Multi-Dompet')
bullet('Buat, edit, hapus wallet dengan warna kustom')
bullet('Saldo di-sync otomatis berdasarkan transaksi terkait')
bullet('Minimal 1 dompet harus ada')

h2('4.3 Budget Bulanan')
bullet('Set budget per kategori per bulan')
bullet('Progress bar hijau → merah saat melebihi budget')
bullet('Perbandingan dengan spending aktual dari categorySummaryProvider')

h2('4.4 Statistik')
bullet('Ringkasan spending per kategori untuk bulan tertentu')
bullet('Progress indicator visual per kategori')
bullet('Navigasi bulan/tahun via selectedMonthProvider & selectedYearProvider')

h2('4.5 Laporan PDF')
bullet('Generate laporan bulanan dalam format PDF')
bullet('Thumbnail lampiran muncul di PDF (gambar max 4 per transaksi, 60x60px)')
bullet('Share via printing package')
bullet('Filename: Catatan_Keuangan_[bulan]_[tahun].pdf')

h2('4.6 Backup & Restore')
bullet('Halaman BackupPage terpisah dengan daftar file backup')
bullet('Metadata: tanggal, ukuran, status encrypted')
bullet('Enkripsi AES-256-CBC dengan PBKDF2 (10,000 iterations)')
bullet('Swipe-to-delete untuk hapus backup')
bullet('Share backup via sistem share sheet')
bullet('Restore dari daftar atau pilih file manual')

h2('4.7 Keamanan')
bullet('PIN Lock: 4–6 digit, setup dengan konfirmasi')
bullet('Biometric: Fingerprint/Face via local_auth')
bullet('Auto-lock: App langsung terkunci saat ke-background (AppLifecycleState.paused)')
bullet('PIN wajib jika biometric aktif')

h2('4.8 Deep Linking')
bullet('Custom scheme: catalog:// atau https://catalog.app/')
bullet('Routes: / (home), /settings, /statistik?tahun=X&bulan=X')
bullet('Error page dengan tombol kembali ke home')

h2('4.9 Android Home Screen Widget')
bullet('Widget menampilkan saldo, pemasukan, pengeluaran bulan ini')
bullet('Update otomatis setiap ada transaksi baru')
bullet('Format ringkas: "1.2jt", "500rb"')
bullet('Green theme, layout Native Android XML')

h2('4.10 Tema')
bullet('Light & Dark mode dengan Material 3')
bullet('Seed color hijau (#2E7D32)')
bullet('Toggle di AppBar atau Settings page')

# ═══════════════════════════════════════════════════════════
# 6. SERVICES & ARSITEKTUR
# ═══════════════════════════════════════════════════════════
h1('5. Services & Arsitektur')

h2('5.1 DatabaseHelper')
body('Singleton class untuk semua operasi SQLite. Menyediakan CRUD lengkap untuk transaksi, dompet, '
     'kategori, budget, dan pengaturan. Menggunakan transaction untuk operasi atomic.')

h2('5.2 CryptoService — AES-256-CBC')
table(
    ['Parameter', 'Nilai'],
    [
        ['Algoritma', 'AES-256-CBC'],
        ['Key Derivation', 'PBKDF2-HMAC-SHA256'],
        ['Iterasi', '10,000'],
        ['Salt', '16 bytes (acak per enkripsi)'],
        ['IV', '16 bytes (acak per enkripsi)'],
        ['Output Encoding', 'Base64'],
        ['Prefix', 'ENCRYPTED:'],
        ['Format', 'ENCRYPTED:<base64(salt+iv+ciphertext)>'],
    ],
    col_widths=[2.0, 4.5]
)

h2('5.3 BiometricService')
body('Wrapper untuk local_auth dengan method: isAvailable(), hasEnrolledBiometrics(), authenticate().')

h2('5.4 FileService')
body('Singleton untuk operasi file: pickImageFromCamera(), pickImageFromGallery(), '
     'pickAnyFile(), pickMultipleFiles(). File disimpan ke folder attachments dengan timestamp.')

h2('5.5 HomeWidgetService')
body('Update Android widget via home_widget package. Menyimpan data saldo ke SharedPreferences '
     'dan memicu update widget provider.')

h2('5.6 PdfLaporanService')
body('Generate PDF laporan bulanan: header, ringkasan total, tabel transaksi per kategori, '
     'dan thumbnail lampiran gambar.')

h2('5.7 State Management — Riverpod 3.x')
body('Semua mutable state menggunakan Notifier class dengan named setter methods.')
code('''// Contoh: SelectedMonthNotifier
class SelectedMonthNotifier extends Notifier<int> {
  @override int build() => DateTime.now().month;
  void increment() { if (state == 12) { state = 1; } else { state++; } }
  void decrement() { if (state == 1) { state = 12; } else { state--; } }
  void setMonth(int month) => state = month;
}''')

# ═══════════════════════════════════════════════════════════
# 7. KEAMANAN & ENKRIPSI
# ═══════════════════════════════════════════════════════════
h1('6. Keamanan & Enkripsi')

h2('6.1 Authentication Flow')
bullet('App startup: Cek apakah PIN sudah dikonfigurasi')
bullet('Jika ya: Tampilkan PinLockScreen')
bullet('Jika biometric enabled: Otomatis coba authenticate')
bullet('Gagal biometric: User masukkan PIN manual')
bullet('App ke-background: _isLocked = true, next resume → PinLockScreen')

h2('6.2 Platform Permissions')
table(
    ['Platform', 'Permission', 'Untuk'],
    [
        ['Android', 'USE_BIOMETRIC', 'Autentikasi fingerprint'],
        ['Android', 'USE_FINGERPRINT', 'Fingerprint legacy'],
        ['Android', 'CAMERA', 'Ambil foto lampiran'],
        ['Android', 'READ/WRITE_EXTERNAL_STORAGE', 'Akses file'],
        ['iOS', 'NSFaceIDUsageDescription', 'Autentikasi Face ID'],
    ],
    col_widths=[1.2, 2.3, 3.0]
)

# ═══════════════════════════════════════════════════════════
# 8. DEEP LINKING & WIDGET
# ═══════════════════════════════════════════════════════════
h1('7. Deep Linking & Home Widget')

h2('7.1 Routes (GoRouter)')
table(
    ['Path', 'Nama', 'Parameter'],
    [
        ['/', 'home', '—'],
        ['/settings', 'settings', '—'],
        ['/statistik', 'statistik', '?tahun=X&bulan=X (opsional)'],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

h2('7.2 Deep Link Schemes')
bullet('Custom scheme: catalog://[path]')
bullet('Contoh: catalog:///statistik?tahun=2026&bulan=4')

h2('7.3 Android Widget Components')
bullet('CatatanKeuanganWidget.kt — AppWidgetProvider')
bullet('widget_catalog_keuangan.xml — Layout XML')
bullet('catalog_keuangan_widget_info.xml — Widget config')
bullet('HomeWidgetService.dart — Flutter service')
bullet('Update flow: App startup → initialize → listen transaksiProvider → '
        'HomeWidgetService.updateWidget() → SharedPreferences → Kotlin render')

# ═══════════════════════════════════════════════════════════
# 9. TESTING & BUILD
# ═══════════════════════════════════════════════════════════
h1('8. Testing & Build')

h2('8.1 Cakupan Test — 63 Test')
table(
    ['File', 'Jumlah', 'Coverage'],
    [
        ['widget_test.dart', '26', 'Serialisasi model, smoke test widget'],
        ['crypto_service_test.dart', '20', 'Roundtrip enkripsi/dekripsi AES, edge cases'],
        ['providers_test.dart', '22', 'Date/month/year notifiers, todayNormalized'],
    ],
    col_widths=[2.5, 1.0, 3.0]
)

h2('8.2 Test Cases CryptoService')
bullet('Roundtrip enkripsi/dekripsi')
bullet('Penanganan password kosong')
bullet('Password salah → FormatException')
bullet('Karakter Unicode dan karakter khusus dalam password')
bullet('Data JSON besar')
bullet('IV acak per enkripsi (ciphertext berbeda setiap kali)')
bullet('Input terenkripsi malformed → FormatException')

h2('8.3 Varian APK')
table(
    ['Variant', 'File', 'Ukuran', 'Target'],
    [
        ['armeabi', 'app-armeabi-release.apk', '~24.7 MB', 'Android ARM v7a (2010+)'],
        ['arm64', 'app-arm64-release.apk', '~26.4 MB', 'Android ARM 64-bit (2015+)'],
        ['universal', 'app-universal-release.apk', '~65.3 MB', 'Semua arsitektur CPU'],
    ],
    col_widths=[1.3, 2.8, 1.2, 1.2]
)

h2('8.4 Build Commands')
code('''# armeabi
flutter build apk --release --target-platform android-arm --no-tree-shake-icons

# arm64
flutter build apk --release --target-platform android-arm64 --no-tree-shake-icons

# universal
flutter build apk --release --no-tree-shake-icons''')

body('Flag --no-tree-shake-icons mencegah korupsi app.so saat build parallel.')

# ── SAVE ──────────────────────────────────────────────────
out = 'C:/Users/muham/catatan_keuangan/Dokumentasi_Catatan_Keuangan_ID.docx'
doc.save(out)
print(f'DISIMPAN: {out}  ({os.path.getsize(out)//1024} KB)')
