"""
Technical Documentation - Personal Finance Tracker
Professional DOCX with green theme styling
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
import datetime
import os

# ── Color Theme (Forest Green) ───────────────────────────
C_DARK     = RGBColor(0x1B, 0x5E, 0x20)   # Deep forest - headings
C_PRIMARY  = RGBColor(0x2E, 0x7D, 0x32)   # Green - accent
C_SECOND   = RGBColor(0x58, 0x81, 0x57)   # Medium green
C_LIGHT    = RGBColor(0xA3, 0xB1, 0x8A)   # Sage - subtle
C_BG       = RGBColor(0xF8, 0xFC, 0xF8)   # Very light green bg
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY     = RGBColor(0x55, 0x55, 0x55)   # Body text
C_LGRAY    = RGBColor(0x88, 0x88, 0x88)   # Light gray
C_TBL_HDR  = RGBColor(0x2E, 0x7D, 0x32)
C_TBL_ALT  = RGBColor(0xF0, 0xF7, 0xF0)   # Alternating row

HEX_DARK   = "1B5E20"
HEX_PRIMARY = "2E7D32"
HEX_SECOND  = "588157"
HEX_LIGHT   = "A3B18A"
HEX_BG      = "F8FCF8"
HEX_WHITE   = "FFFFFF"
HEX_GRAY    = "555555"
HEX_LGRAY   = "888888"
HEX_ALT     = "F0F7F0"

# ── Helpers ──────────────────────────────────────────────
doc = None

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def para_spacing(para, before=0, after=6):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before * 20))
    spacing.set(qn('w:after'), str(after * 20))
    pPr.append(spacing)

def set_run_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

# ── Document Setup ────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = C_GRAY
pPr = style._element.get_or_add_pPr()
spacing = OxmlElement('w:spacing')
spacing.set(qn('w:after'), '80')
pPr.append(spacing)

# ── COVER PAGE ──────────────────────────────────────────────
# Top accent bar
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run(' ')
run.font.size = Pt(6)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '24')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), HEX_PRIMARY)
pBdr.append(bottom)
pPr.append(pBdr)

# Spacer
for _ in range(2):
    doc.add_paragraph()

# App name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p, 0, 4)
run = p.add_run('PERSONAL FINANCE TRACKER')
run.font.name = 'Calibri'
run.font.size = Pt(42)
run.font.bold = True
run.font.color.rgb = C_PRIMARY

# Tagline
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p, 0, 24)
run = p.add_run('Personal Finance Tracker')
run.font.name = 'Calibri'
run.font.size = Pt(20)
run.font.color.rgb = C_SECOND

# Divider line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p, 0, 24)
run = p.add_run('─' * 50)
run.font.color.rgb = C_LIGHT

# Info block
info = [
    ('Version', '1.0.0'),
    ('Platform', 'Flutter / Android / iOS'),
    ('Date', datetime.date.today().strftime('%d %B %Y')),
    ('Language', 'Dart 3.11 / Flutter 3.41'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, 0, 4)
    run = p.add_run(f'{label}: ')
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = C_LGRAY
    run2 = p.add_run(value)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(12)
    run2.font.color.rgb = C_GRAY

# Bottom bar
for _ in range(1):
    doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run(' ')
run.font.size = Pt(4)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single')
top.set(qn('w:sz'), '24')
top.set(qn('w:space'), '1')
top.set(qn('w:color'), HEX_PRIMARY)
pBdr.append(top)
pPr.append(pBdr)

doc.add_page_break()

# ── Helper: Section Heading ────────────────────────────────
def h1(text):
    # Green bar + H1 text
    p = doc.add_paragraph()
    para_spacing(p, 20, 8)
    # Left accent bar via border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '36')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), HEX_PRIMARY)
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = C_DARK
    return p

def h2(text):
    p = doc.add_paragraph()
    para_spacing(p, 16, 4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = C_PRIMARY
    return p

def h3(text):
    p = doc.add_paragraph()
    para_spacing(p, 10, 2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = C_SECOND
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    para_spacing(p, 2, 6)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = C_GRAY
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    para_spacing(p, 1, 3)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = C_GRAY
    return p

def code(text):
    p = doc.add_paragraph()
    para_spacing(p, 4, 4)
    p.paragraph_format.left_indent = Inches(0.3)
    # Light gray bg
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4F0')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = C_DARK
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, HEX_PRIMARY)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = C_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            set_cell_bg(cell, HEX_ALT if ri % 2 == 0 else HEX_WHITE)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = C_GRAY

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return t

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════
h1('Table of Contents')

toc_items = [
    ('1', 'Overview', '3'),
    ('2', 'Tech Stack & Structure', '4'),
    ('3', 'Database Schema', '4'),
    ('4', 'Main Features', '5'),
    ('5', 'Services & Architecture', '7'),
    ('6', 'Security & Encryption', '8'),
    ('7', 'Deep Linking & Home Widget', '9'),
    ('8', 'Testing & Build', '10'),
]

for num, title, pg in toc_items:
    p = doc.add_paragraph()
    para_spacing(p, 3, 3)
    run = p.add_run(f'{num}.  ')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = C_PRIMARY
    run2 = p.add_run(title)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(11)
    run2.font.color.rgb = C_GRAY
    run3 = p.add_run(f'  ...............  {pg}')
    run3.font.name = 'Calibri'
    run3.font.size = Pt(10)
    run3.font.color.rgb = C_LIGHT

page_break()

# ═══════════════════════════════════════════════════════════
# 2. OVERVIEW
# ═══════════════════════════════════════════════════════════
h1('1. Overview')

body('Personal Finance Tracker is a personal financial record-keeping application developed with Flutter. '
     'Designed to help users manage personal finances with comprehensive features including transaction '
     'recording, multi-wallet, monthly budgets, recurring transactions, PDF reports with attachments, '
     'AES-256 encrypted backup, biometric authentication, and Android home screen widget.')

h2('1.1 Project Information')
table(
    ['Property', 'Value'],
    [
        ['Project Name', 'catatan_keuangan'],
        ['Package', 'com.example.catatan_keuangan'],
        ['Version', '1.0.0'],
        ['Flutter SDK', '3.41.x'],
        ['Dart SDK', '3.11.x'],
        ['Target Platform', 'Android (minSdk 21), iOS'],
        ['State Management', 'Riverpod 3.x'],
        ['Routing', 'GoRouter 14.x'],
    ],
    col_widths=[2.0, 4.5]
)

h2('1.2 Code Statistics')
table(
    ['Metric', 'Value'],
    [
        ['Total code lines (lib/)', '~2,800 lines'],
        ['Total Dart files', '23 files'],
        ['Test cases', '63 tests (3 test files)'],
        ['Dependencies', '15 packages'],
        ['Database version', 'v8 (with migration)'],
        ['APK variants', '3 (armeabi, arm64, universal)'],
    ],
    col_widths=[3.0, 3.5]
)

# ═══════════════════════════════════════════════════════════
# 3. TECH STACK & STRUCTURE
# ═══════════════════════════════════════════════════════════
h1('2. Tech Stack & Structure')

h2('2.1 Dependencies')
table(
    ['Category', 'Package', 'Version', 'Function'],
    [
        ['Database', 'sqflite', '^2.3.0', 'SQLite ORM'],
        ['State', 'flutter_riverpod', '^3.3.1', 'State management'],
        ['Navigation', 'go_router', '^14.8.1', 'Routing + deep linking'],
        ['PDF', 'pdf + printing', '3.10.8 / 5.12.0', 'Generate & share PDF'],
        ['File', 'image_picker + file_picker', 'various', 'Camera, gallery, file picker'],
        ['Encryption', 'encrypt + pointycastle', '5.0.3 / 3.9.1', 'AES-256-CBC encryption'],
        ['Biometric', 'local_auth', '^2.3.0', 'Fingerprint / Face ID'],
        ['Widget', 'home_widget', '^0.6.0', 'Android home widget'],
        ['Format', 'intl', '^0.20.2', 'Indonesian locale'],
    ],
    col_widths=[1.2, 1.8, 1.3, 2.2]
)

h2('2.2 Directory Structure')
code('''lib/
├── main.dart                  Entry point
├── app.dart                   MaterialApp + lock overlay
├── router.dart                 GoRouter + deep linking
├── providers.dart              Riverpod providers
├── models/
│   ├── models.dart            Transaksi, Dompet, Kategori, Budget, Pengaturan
│   └── constants.dart         Application constants
├── services/
│   ├── database_helper.dart   SQLite CRUD + migration (v8)
│   ├── crypto_service.dart     AES-256-CBC + PBKDF2
│   ├── biometric_service.dart  local_auth wrapper
│   ├── file_service.dart       Image/file picking
│   ├── home_widget_service.dart Android widget updater
│   └── pdf_laporan_service.dart PDF generator
└── pages/
    ├── home_page.dart          TabBar (Hari Ini / Per Tanggal / Bulanan / Lainnya)
    ├── settings_page.dart      Dark mode + biometric toggle
    ├── statistik_page.dart      Spending statistics
    ├── backup_page.dart        Backup & Restore UI
    ├── pin_lock_screen.dart    PIN + biometric auth
    ├── sheets/                 Bottom sheet components
    └── widgets/                 Reusable widget components''')

# ═══════════════════════════════════════════════════════════
# 4. DATABASE SCHEMA
# ═══════════════════════════════════════════════════════════
h1('3. Database Schema')

body('SQLite with version v8 and auto-migration system. 6 indexes for query performance.')

h2('3.1 Database Tables')
table(
    ['Table', 'Main Columns', 'Description'],
    [
        ['transaksi', 'id, jenis, jumlah, deskripsi, kategori,\ntanggal, lampiran, is_recurring,\nrecurring_frequency, id_dompet, deleted_at',
         'Transactions with soft delete'],
        ['dompet', 'id, nama, saldo, warna', 'Multi-wallet with balance'],
        ['kategori', 'id, nama, jenis, icon, is_default', 'Categories with icon'],
        ['budget', 'id, bulan, tahun, nominal, kategori', 'Monthly budget per category'],
        ['pengaturan', 'id, is_dark_mode, pin, use_biometric', 'Application settings'],
    ],
    col_widths=[1.2, 3.3, 2.0]
)

h2('3.2 Database Indexes')
for idx in ['idx_transaksi_tanggal', 'idx_transaksi_kategori', 'idx_transaksi_jenis',
           'idx_transaksi_dompet', 'idx_budget_period', 'idx_budget_unique (UNIQUE)']:
    bullet(f'{idx}')

h2('3.3 Migration History')
table(
    ['Version', 'Changes'],
    [
        ['v1–v3', 'Added dompet, budget, pengaturan tables'],
        ['v4', 'Added lampiran column to transaksi'],
        ['v5', 'Added kategori table'],
        ['v6', 'Added indexes for query performance'],
        ['v7', 'Added soft delete (deleted_at column)'],
        ['v8', 'Added use_biometric to pengaturan'],
    ],
    col_widths=[1.0, 5.5]
)

# ═══════════════════════════════════════════════════════════
# 5. MAIN FEATURES
# ═══════════════════════════════════════════════════════════
h1('4. Main Features')

h2('4.1 Transaction Recording')
body('Core application feature for recording transactions.')
bullet('Type: Income or Expense')
bullet('Category: 9 defaults for expenses + 6 for income, custom with icon')
bullet('Attachment: Image from camera/gallery or arbitrary file')
bullet('Recurring: Daily, Weekly, Monthly, Yearly — auto-processed on app start')
bullet('Soft delete: Swipe left to delete, moved to trash, can be restored')
bullet('Tap to edit, dismiss to delete')

h2('4.2 Multi-Wallet')
bullet('Create, edit, delete wallet with custom color')
bullet('Balance auto-synced based on related transactions')
bullet('At least 1 wallet must exist')

h2('4.3 Monthly Budget')
bullet('Set budget per category per month')
bullet('Progress bar green → red when exceeding budget')
bullet('Comparison with actual spending from categorySummaryProvider')

h2('4.4 Statistics')
bullet('Spending summary per category for a specific month')
bullet('Visual progress indicator per category')
bullet('Month/year navigation via selectedMonthProvider & selectedYearProvider')

h2('4.5 PDF Reports')
bullet('Generate monthly reports in PDF format')
bullet('Attachment thumbnails shown in PDF (max 4 images per transaction, 60x60px)')
bullet('Share via printing package')
bullet('Filename: Catatan_Keuangan_[month]_[year].pdf')

h2('4.6 Backup & Restore')
bullet('Separate BackupPage with list of backup files')
bullet('Metadata: date, size, encrypted status')
bullet('AES-256-CBC encryption with PBKDF2 (10,000 iterations)')
bullet('Swipe-to-delete to delete backup')
bullet('Share backup via system share sheet')
bullet('Restore from list or select file manually')

h2('4.7 Security')
bullet('PIN Lock: 4–6 digits, setup with confirmation')
bullet('Biometric: Fingerprint/Face via local_auth')
bullet('Auto-lock: App immediately locked when going to background (AppLifecycleState.paused)')
bullet('PIN required if biometric is active')

h2('4.8 Deep Linking')
bullet('Custom scheme: catalog:// or https://catalog.app/')
bullet('Routes: / (home), /settings, /statistik?tahun=X&bulan=X')
bullet('Error page with back-to-home button')

h2('4.9 Android Home Screen Widget')
bullet('Widget displays balance, income, expense for current month')
bullet('Auto-update on every new transaction')
bullet('Compact format: "1.2jt", "500rb"')
bullet('Green theme, Native Android XML layout')

h2('4.10 Theme')
bullet('Light & Dark mode with Material 3')
bullet('Seed color green (#2E7D32)')
bullet('Toggle in AppBar or Settings page')

# ═══════════════════════════════════════════════════════════
# 6. SERVICES & ARCHITECTURE
# ═══════════════════════════════════════════════════════════
h1('5. Services & Architecture')

h2('5.1 DatabaseHelper')
body('Singleton class for all SQLite operations. Provides complete CRUD for transactions, dompet, '
     'categories, budget, and settings. Uses transactions for atomic operations.')

h2('5.2 CryptoService — AES-256-CBC')
table(
    ['Parameter', 'Value'],
    [
        ['Algorithm', 'AES-256-CBC'],
        ['Key Derivation', 'PBKDF2-HMAC-SHA256'],
        ['Iterations', '10,000'],
        ['Salt', '16 bytes (random per encryption)'],
        ['IV', '16 bytes (random per encryption)'],
        ['Output Encoding', 'Base64'],
        ['Prefix', 'ENCRYPTED:'],
        ['Format', 'ENCRYPTED:<base64(salt+iv+ciphertext)>'],
    ],
    col_widths=[2.0, 4.5]
)

h2('5.3 BiometricService')
body('Wrapper for local_auth with methods: isAvailable(), hasEnrolledBiometrics(), authenticate().')

h2('5.4 FileService')
body('Singleton for file operations: pickImageFromCamera(), pickImageFromGallery(), '
     'pickAnyFile(), pickMultipleFiles(). Files saved to attachments folder with timestamp.')

h2('5.5 HomeWidgetService')
body('Update Android widget via home_widget package. Stores balance data in SharedPreferences '
     'and triggers widget provider update.')

h2('5.6 PdfLaporanService')
body('Generate monthly PDF reports: header, total summary, transaction table per category, '
     'and image attachment thumbnails.')

h2('5.7 State Management — Riverpod 3.x')
body('All mutable state uses Notifier class with named setter methods.')
code('''// Contoh: SelectedMonthNotifier
class SelectedMonthNotifier extends Notifier<int> {
  @override int build() => DateTime.now().month;
  void increment() { if (state == 12) { state = 1; } else { state++; } }
  void decrement() { if (state == 1) { state = 12; } else { state--; } }
  void setMonth(int month) => state = month;
}''')

# ═══════════════════════════════════════════════════════════
# 7. SECURITY & ENCRYPTION
# ═══════════════════════════════════════════════════════════
h1('6. Security & Encryption')

h2('6.1 Authentication Flow')
bullet('App startup: Check if PIN is configured')
bullet('If yes: Show PinLockScreen')
bullet('If biometric enabled: Auto attempt authenticate')
bullet('Biometric failed: User enters PIN manually')
bullet('App to background: _isLocked = true, next resume → PinLockScreen')

h2('6.2 Platform Permissions')
table(
    ['Platform', 'Permission', 'For'],
    [
        ['Android', 'USE_BIOMETRIC', 'Fingerprint authentication'],
        ['Android', 'USE_FINGERPRINT', 'Legacy fingerprint'],
        ['Android', 'CAMERA', 'Capture photo attachments'],
        ['Android', 'READ/WRITE_EXTERNAL_STORAGE', 'File access'],
        ['iOS', 'NSFaceIDUsageDescription', 'Face ID authentication'],
    ],
    col_widths=[1.2, 2.3, 3.0]
)

# ═══════════════════════════════════════════════════════════
# 8. DEEP LINKING & WIDGET
# ═══════════════════════════════════════════════════════════
h1('7. Deep Linking & Home Widget')

h2('7.1 Routes (GoRouter)')
table(
    ['Path', 'Name', 'Parameters'],
    [
        ['/', 'home', '—'],
        ['/settings', 'settings', '—'],
        ['/statistik', 'statistik', '?tahun=X&bulan=X (optional)'],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

h2('7.2 Deep Link Schemes')
bullet('Custom scheme: catalog://[path]')
bullet('Example: catalog:///statistik?tahun=2026&bulan=4')

h2('7.3 Android Widget Components')
bullet('CatatanKeuanganWidget.kt — AppWidgetProvider')
bullet('widget_catalog_keuangan.xml — Layout XML')
bullet('catalog_keuangan_widget_info.xml — Widget config')
bullet('HomeWidgetService.dart — Flutter service')
bullet('Update flow: App startup → initialize → listen transaksiProvider → '
        'HomeWidgetService.updateWidget() → SharedPreferences → Kotlin render')

# ═══════════════════════════════════════════════════════════
# 9. TESTING & BUILD
# ═══════════════════════════════════════════════════════════
h1('8. Testing & Build')

h2('8.1 Test Coverage — 63 Tests')
table(
    ['File', 'Count', 'Coverage'],
    [
        ['widget_test.dart', '26', 'Model serialization, widget smoke test'],
        ['crypto_service_test.dart', '20', 'AES encrypt/decrypt roundtrip, edge cases'],
        ['providers_test.dart', '22', 'Date/month/year notifiers, todayNormalized'],
    ],
    col_widths=[2.5, 1.0, 3.0]
)

h2('8.2 CryptoService Test Cases')
bullet('encrypt/decrypt roundtrip')
bullet('Empty password handling')
bullet('Wrong password → FormatException')
bullet('Unicode characters & special characters in password')
bullet('Large JSON data')
bullet('Random IV per encryption (different ciphertext each time)')
bullet('Malformed encrypted input → FormatException')

h2('8.3 APK Variants')
table(
    ['Variant', 'File', 'Size', 'Target'],
    [
        ['armeabi', 'app-armeabi-release.apk', '~24.7 MB', 'Android ARM v7a (2010+)'],
        ['arm64', 'app-arm64-release.apk', '~26.4 MB', 'Android ARM 64-bit (2015+)'],
        ['universal', 'app-universal-release.apk', '~65.3 MB', 'All architectures'],
    ],
    col_widths=[1.3, 2.8, 1.2, 1.2]
)

h2('8.4 Build Commands')
code('''# armeabi
flutter build apk --release --target-platform android-arm --no-tree-shake-icons

# arm64
flutter build apk --release --target-platform android-arm64 --no-tree-shake-icons

# universal
flutter build apk --release --no-tree-shake-icons''')

body('Flag --no-tree-shake-icons prevents app.so corruption during parallel builds.')

# ── SAVE ──────────────────────────────────────────────────
out = 'C:/Users/muham/catatan_keuangan/Dokumentasi_Catatan_Keuangan_EN.docx'
doc.save(out)
print(f'SAVED: {out}  ({os.path.getsize(out)//1024} KB)')
