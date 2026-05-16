from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Page Setup ───────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.8)
section.bottom_margin = Inches(0.8)

# ─── Colors ────────────────────────────────────────────────────
GREEN_DARK  = RGBColor(0x1B, 0x5E, 0x20)
GREEN_MID   = RGBColor(0x2E, 0x7D, 0x32)
GREEN_LIGHT = RGBColor(0x4C, 0xAF, 0x50)
CORAL       = RGBColor(0xE5, 0x73, 0x73)
GOLD        = RGBColor(0xFF, 0xB3, 0x00)
DARK_TEXT   = RGBColor(0x1A, 0x1A, 0x2E)
GRAY_TEXT   = RGBColor(0x6B, 0x72, 0x80)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# ─── Helper: set paragraph shading ─────────────────────────────
def shade_paragraph(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    pPr.append(shd)

# ─── Helper: add colored heading ─────────────────────────────
def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15 if level == 1 else 13)
    if level == 1:
        run.font.color.rgb = GREEN_DARK
    else:
        run.font.color.rgb = GREEN_MID
    return p

# ─── Helper: add bullet ─────────────────────────────────────
def add_bullet(doc, text, sub=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent = Inches(0.3 if sub else 0.0)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK_TEXT
    return p

# ─── Helper: add normal paragraph ─────────────────────────────
def add_para(doc, text, bold=False, color=None, size=10, space_before=2, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color or DARK_TEXT
    return p

# ─── Helper: styled table ─────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1B5E20')
        tcPr.append(shd)

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        fill = 'F1F8F1' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9)
            cells[c_idx].paragraphs[0].runs[0].font.color.rgb = DARK_TEXT
            tc = cells[c_idx]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  fill)
            tcPr.append(shd)

    if col_widths:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]

    doc.add_paragraph()
    return tbl

# ═══════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('DompetKu')
run.font.size  = Pt(40)
run.bold       = True
run.font.color.rgb = GREEN_DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Aplikasi Pencatatan Keuangan Personal')
run2.font.size  = Pt(16)
run2.font.color.rgb = GRAY_TEXT

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Spesifikasi Teknis & Dokumentasi Fitur')
run3.font.size  = Pt(13)
run3.font.color.rgb = GREEN_MID

doc.add_paragraph()

# Version badge
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p4, '1B5E20')
run4 = p4.add_run('  v1.0.0 — April 2026  ')
run4.font.size  = Pt(12)
run4.bold       = True
run4.font.color.rgb = WHITE

doc.add_paragraph()
doc.add_paragraph()

# Tech stack
add_table(doc,
    ['Teknologi', 'Detail'],
    [
        ['Framework',    'Flutter'],
        ['State Mgmt',  'Riverpod (flutter_riverpod)'],
        ['Database',     'SQLite (sqflite)'],
        ['Navigation',   'GoRouter'],
        ['PDF Gen',     'pdf + printing'],
        ['Charts',      'fl_chart'],
        ['Animations',  'Lottie'],
        ['Localization', 'flutter_localizations (ID + EN)'],
        ['Build Target','Android arm64-v8a (release)'],
    ],
    col_widths=[Inches(2.2), Inches(4.0)]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 1: DATABASE
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '1. DATABASE', 1)
add_para(doc, 'Aplikasi menggunakan SQLite sebagai primary database via package sqflite. Schema mendukung migrasi dari versi 1 hingga versi 10.', size=9, color=GRAY_TEXT)

doc.add_paragraph()
add_section_heading(doc, 'Tabel Database', 2)
add_table(doc,
    ['Tabel', 'Deskripsi', 'Key Columns'],
    [
        ['transaksi',        'Pemasukan & Pengeluaran',           'id, jenis, jumlah, tanggal, kategori'],
        ['dompet',           'Multi-wallet',                     'id, nama, saldo, warna, currency'],
        ['kategori',         'Kategori transaksi',                'id, nama, jenis, icon, is_default'],
        ['budget',           'Anggaran bulanan per kategori',     'id, bulan, tahun, nominal, kategori'],
        ['pengaturan',       'App settings (dark mode, PIN)',   'id, is_dark_mode, pin, use_biometric'],
        ['utang_piutang',    'Catatan utang & piutang',         'id, nama_orang, jenis, nominal_total'],
        ['history_cicilan',  'Riwayat pembayaran cicilan',        'id, id_utang_piutang, nominal, tanggal'],
        ['tabungan_impian',   'Target tabungan impian',            'id, nama_impian, target_nominal, terkumpul'],
    ],
    col_widths=[Inches(1.6), Inches(2.8), Inches(2.8)]
)

add_para(doc, 'Indeks: tanggal, kategori, jenis, dompet, budget_period. Semua kolom unik di-handle via UNIQUE constraint.', size=8.5, color=GRAY_TEXT, space_before=6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 2: FITUR TRANSAKSI
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '2. TRANSAKSI', 1)

add_section_heading(doc, '2.1 Tambah Transaksi (tambah_transaksi_sheet.dart)', 2)
add_bullet(doc, 'Toggle jenis transaksi: Pemasukan / Pengeluaran')
add_bullet(doc, 'Date picker dengan batasan minYear (2020) dan maxFutureDays (365)')
add_bullet(doc, 'Input nominal dengan CurrencyInputFormatter otomatis format Rupiah Indonesia')
add_bullet(doc, 'Selector dompet (multi-wallet support)')
add_bullet(doc, 'Deskripsi — OPSIONAL (empty string jika kosong)')
add_bullet(doc, 'Kategori chips berbasis kategori provider (dinamis dari database)')
add_bullet(doc, 'Toggle transaksi berulang: Harian / Mingguan / Bulanan / Triwulanan / Tahunan')
add_bullet(doc, 'Lampiran multi-file via camera, gallery, atau file picker')
add_bullet(doc, 'OCR scan struk via AI (OCRService) — auto-read nominal dari foto')
add_bullet(doc, 'Soft delete via swipe gesture (Dismissible)')
add_bullet(doc, 'Auto-sync dompet saldo setelah insert/update/delete transaksi')

add_section_heading(doc, '2.2 Navigasi Transaksi', 2)
add_table(doc,
    ['Tab', 'File', 'Fitur'],
    [
        ['Hari Ini',    'tab_hari_ini.dart',   'Transaksi hari ini, header premium gradient, animated counters'],
        ['Per Tanggal',  'tab_per_tanggal.dart','Kalender interaktif TableCalendar, markers dots, paginated list'],
        ['Bulanan',      'tab_bulanan.dart',    'Navigasi bulan, ringkasan 3 kartu, budget warning, paginated list'],
        ['Lainnya',      'tab_lainnya.dart',    'Grid menu ke semua fitur, panduan aplikasi'],
        ['Dashboard',    'tab_dashboard.dart',  'Balance, income/expense, budget, weekly chart, pie chart'],
    ],
    col_widths=[Inches(1.1), Inches(1.8), Inches(4.3)]
)

add_section_heading(doc, '2.3 Pencarian Transaksi (transaksi_search_delegate.dart)', 2)
add_bullet(doc, 'Full-text search via SearchDelegate')
add_bullet(doc, 'Mencari di kolom deskripsi dan kategori')
add_bullet(doc, 'Live FutureBuilder results')
add_bullet(doc, 'Tap hasil untuk langsung edit transaksi')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 3: DOMPET & KATEGORI
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '3. DOMPET & KATEGORI', 1)

add_section_heading(doc, '3.1 Kelola Dompet (kelola_dompet_sheet.dart)', 2)
add_bullet(doc, 'Tambah dompet baru: nama + warna (8 pilihan warna)')
add_bullet(doc, 'Edit nama dan warna dompet')
add_bullet(doc, 'Hapus dompet dengan konfirmasi dialog')
add_bullet(doc, 'Saldo auto-sync dari transaksi terkait via syncDompetSaldo()')

add_section_heading(doc, '3.2 Kelola Kategori (kelola_kategori_sheet.dart)', 2)
add_bullet(doc, '12 kategori pengeluaran default: Makanan, Transportasi, Belanja, Hiburan, Tagihan, Kesehatan, Rumah Tangga, Pendidikan, Fashion, Pulsa & Data, Usaha/Bisnis, Lainnya')
add_bullet(doc, '6 kategori pemasukan default: Gaji, Bonus, Usaha, Investasi, Hadiah, Lainnya')
add_bullet(doc, 'Tambah/edit/hapus kategori kustom')
add_bullet(doc, 'Ikon per kategori (Material Icons)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 4: BUDGET & TABUNGAN
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '4. BUDGET & TABUNGAN', 1)

add_section_heading(doc, '4.1 Budget Bulanan (budget_sheet.dart)', 2)
add_bullet(doc, 'Set budget per kategori per bulan')
add_bullet(doc, 'UNIQUE constraint: satu budget per kategori per bulan')
add_bullet(doc, 'Progress bar visual per kategori')
add_bullet(doc, 'Warning saat budget ≥80% terpakai')
add_bullet(doc, 'Budget warning chip di Tab Bulanan')
add_bullet(doc, 'Dialog inline untuk edit nominal budget')

add_section_heading(doc, '4.2 Tabungan Impian (tabungan_impian_page.dart)', 2)
add_bullet(doc, 'Grid card dengan CircularProgressIndicator')
add_bullet(doc, 'Persentase dan nominal terkumpul dari target')
add_bullet(doc, 'Tambah target baru: nama, nominal, tanggal tenggat')
add_bullet(doc, 'Tabung progress — kurangi saldo dompet + buat transaksi pengeluaran')
add_bullet(doc, 'Sync tabungan ke transaksi pengeluaran secara otomatis')
add_bullet(doc, 'Hapus target tabungan')
add_bullet(doc, 'Trigger achievement "Sang Visionaris" saat ≥2 target dibuat')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 5: UTANG PIUTANG
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '5. UTANG PIUTANG', 1)

add_section_heading(doc, 'utang_piutang_page.dart', 2)
add_bullet(doc, 'Tab "Saya Berutang" (utang = kita berutang ke orang lain)')
add_bullet(doc, 'Tab "Orang Berutang" (piutang = orang berutang ke kita)')
add_bullet(doc, 'Card dengan nama, nominal, progres cicilan, tanggal, tenggat')
add_bullet(doc, 'Progress bar visual cicilan')
add_bullet(doc, 'Tambah utang/piutang baru')
add_bullet(doc, 'Cicilan parsial — bayar sebagian nominal')
add_bullet(doc, 'Tandai lunas / belum lunas')
add_bullet(doc, 'History cicilan per utang/piutang')
add_bullet(doc, 'Notification reminder saat jatuh tempo')
add_bullet(doc, 'Achievement "Tepat Janji" saat lunasi utang pertama')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 6: LAPORAN PDF
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '6. LAPORAN PDF', 1)

add_section_heading(doc, '6.1 Generate PDF (pdf_laporan_service.dart)', 2)
add_bullet(doc, 'Pilih periode (tahun + bulan) dan dompet (spesifik atau semua)')
add_bullet(doc, 'Summary card: Total Pemasukan, Total Pengeluaran, Saldo')
add_bullet(doc, 'Pie chart breakdown Pengeluaran per Kategori (SUM query)')
add_bullet(doc, 'Pie chart breakdown Pemasukan per Kategori (SUM query)')
add_bullet(doc, 'Daftar transaksi grouped by date dengan day total')
add_bullet(doc, 'Lampiran gambar di dalam transaksi')
add_bullet(doc, 'Footer: nomor halaman otomatis')
add_bullet(doc, 'Print langsung atau share/export via printing package')

add_section_heading(doc, '6.2 Laporan Sheet (laporan_sheet.dart)', 2)
add_bullet(doc, 'ChoiceChip untuk tahun (5 tahun terakhir)')
add_bullet(doc, 'ChoiceChip untuk bulan (12 bulan)')
add_bullet(doc, 'ChoiceChip untuk dompet (semua + list dompet)')
add_bullet(doc, 'Preview periode yang dipilih')
add_bullet(doc, 'Tombol "Bagikan" (share) dan "Cetak/Simpan" (print)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 7: DASHBOARD
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '7. DASHBOARD', 1)
add_table(doc,
    ['Widget', 'File', 'Fitur'],
    [
        ['Balance Overview',      'balance_overview_card.dart',   'Total saldo, animated counter, dompet aktif count, state persistence'],
        ['Income/Expense Row',   'income_expense_row.dart',      'Pemasukan & Pengeluaran bulanan, animated gradient cards'],
        ['Budget Progress',       'budget_progress_section.dart',  'Progress bar per kategori, warning chips ≥80%'],
        ['Savings Goals',         'savings_goals_section.dart',    'Top 3 tabungan aktif, progress bar'],
        ['Debt Summary',          'debt_summary_section.dart',      'Total utang & piutang aktif, count badge'],
        ['Weekly Spending Chart', 'weekly_spending_chart.dart',     'Bar chart 7 hari terakhir'],
        ['Category Pie Chart',    'category_pie_chart.dart',       'Pie chart pengeluaran per kategori, legend'],
    ],
    col_widths=[Inches(1.8), Inches(2.2), Inches(3.2)]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 8: BACKUP, EXPORT, STATISTIK
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '8. BACKUP, EXPORT & STATISTIK', 1)

add_section_heading(doc, '8.1 Backup & Restore (backup_restore_page.dart)', 2)
add_bullet(doc, 'Backup JSON: ekspor semua 8 tabel ke file JSON')
add_bullet(doc, 'Restore: impor dari file JSON backup, konfirmasi sebelum overwrite')
add_bullet(doc, 'Timestamp backup terakhir (SharedPreferences)')
add_bullet(doc, 'Semua provider di-invalidate + signal broadcast setelah restore')
add_bullet(doc, 'Navigasi ke home setelah restore selesai')

add_section_heading(doc, '8.2 Export CSV (export_service.dart)', 2)
add_bullet(doc, 'Ekspor transaksi ke file CSV (Excel-compatible)')
add_bullet(doc, 'UTF-8 BOM prefix untuk kompatibilitas Excel')
add_bullet(doc, 'Delimiter: semicolon (;) dengan escape quotes')

add_section_heading(doc, '8.3 Statistik (statistik_page.dart)', 2)
add_bullet(doc, 'Pie chart pengeluaran per kategori (fl_chart)')
add_bullet(doc, 'Gradient header dengan bulan berjalan')
add_bullet(doc, 'Total pengeluaran animated counter')
add_bullet(doc, 'Sorted list per kategori dengan progress bar')
add_bullet(doc, 'Lottie empty state animation')
add_bullet(doc, 'Parameter tahun/bulan dari URL query (deep link support)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 9: TRANSAKSI BERULANG & TONG SAMPAH
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '9. TRANSAKSI BERULANG & TONG SAMPAH', 1)

add_section_heading(doc, '9.1 Transaksi Berulang (recurring_transaksi_sheet.dart + recurring_scheduler.dart)', 2)
add_bullet(doc, 'List semua recurring transactions')
add_bullet(doc, 'Frekuensi: harian, mingguan, bulanan, triwulanan, tahunan')
add_bullet(doc, 'Scheduler auto-create: cek sekali per hari via SharedPreferences')
add_bullet(doc, 'Auto-create transaksi saat app di-resume atau di-start')
add_bullet(doc, 'Notification saat transaksi auto-dibuat')
add_bullet(doc, 'Signal broadcast ke transaksi provider setelah auto-create')

add_section_heading(doc, '9.2 Tong Sampah (trash_sheet.dart)', 2)
add_bullet(doc, 'Soft delete transactions (deleted_at timestamp)')
add_bullet(doc, 'List transaksi yang dihapus dengan timestamp deletion')
add_bullet(doc, 'Restore transaksi — broadcast signal + invalidate providers')
add_bullet(doc, 'Hapus permanen dengan konfirmasi dialog')
add_bullet(doc, 'Informasi "30 hari" (soft constraint)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 10: PENGATURAN & PENGAMANAN
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '10. PENGATURAN & PENGAMANAN', 1)

add_section_heading(doc, '10.1 Settings (settings_page.dart)', 2)
add_bullet(doc, 'Toggle Dark Mode — realtime switch, persisted ke database')
add_bullet(doc, 'PIN Lock setup dan verifikasi (4 digit)')
add_bullet(doc, 'Biometric unlock (fingerprint/face) toggle')
add_bullet(doc, 'Notification switches: Budget, Utang, Tabungan')
add_bullet(doc, 'Backup & Restore shortcut link')
add_bullet(doc, 'About section dengan versi aplikasi')

add_section_heading(doc, '10.2 PIN Lock (pin_lock_screen.dart)', 2)
add_bullet(doc, '4-digit PIN via Number Pad UI')
add_bullet(doc, 'Setup baru atau verifikasi PIN')
add_bullet(doc, 'Auto-lock saat app masuk state paused')
add_bullet(doc, 'Biometric unlock jika enabled di settings')

add_section_heading(doc, '10.3 Biometric Service (biometric_service.dart)', 2)
add_bullet(doc, 'Check device biometric availability')
add_bullet(doc, 'Check enrolled biometrics')
add_bullet(doc, 'Authenticate wrapper dengan error handling')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 11: NOTIFICATIONS & ACHIEVEMENTS
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '11. NOTIFICATIONS & ACHIEVEMENTS', 1)

add_section_heading(doc, '11.1 Notification Service (notification_service.dart)', 2)
add_bullet(doc, '5 channel: Budget, Debt, Savings, Achievement, Recurring')
add_bullet(doc, 'Budget warning saat ≥80% budget terpakai')
add_bullet(doc, 'Debt due reminder saat jatuh tempo atau overdue')
add_bullet(doc, 'Savings milestone notification (≥50%, 100%)')
add_bullet(doc, 'Achievement unlock notification')
add_bullet(doc, 'Recurring auto-create notification')
add_bullet(doc, 'Deduplication per session via in-memory Set')
add_bullet(doc, 'Permission request Android 13+')
add_bullet(doc, 'Auto-check on app start (checkAndNotify)')

add_section_heading(doc, '11.2 Pencapaian / Lencana (achievement_page.dart)', 2)
add_table(doc,
    ['ID', 'Nama', 'Deskripsi', 'Syarat'],
    [
        ['first_step',   'Langkah Pertama',    'Mencatat transaksi pertama kali',         '≥1 transaksi'],
        ['consistent',   'Si Konsisten',       'Mencatat lebih dari 50 transaksi',      '≥50 transaksi'],
        ['sultan',       'Sultan',             'Total pendapatan ≥Rp10 Juta',           'totalIn ≥ 10jt'],
        ['responsible',  'Tepat Janji',       'Melunasi utang untuk pertama kali',      'lunasCount ≥ 1'],
        ['visionary',    'Sang Visionaris',     'Membuka minimal 2 tabungan impian',     'savingsCount ≥ 2'],
        ['big_boss',     'Bos Besar',           'Mencatat lebih dari 500 transaksi',     '≥500 transaksi'],
    ],
    col_widths=[Inches(1.0), Inches(1.4), Inches(2.4), Inches(2.4)]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 12: STATE MANAGEMENT
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '12. STATE MANAGEMENT (Riverpod)', 1)

add_section_heading(doc, '12.1 Provider Architecture', 2)
add_table(doc,
    ['Provider', 'Tipe', 'Dependency'],
    [
        ['todayNormalizedProvider',         'Provider (stable)',       'DateTime.now() captured once'],
        ['selectedViewDateProvider',        'NotifierProvider',        'DateTime (nullable)'],
        ['selectedMonthProvider',           'NotifierProvider',        'int (bulan)'],
        ['selectedYearProvider',             'NotifierProvider',        'int (tahun)'],
        ['updateSignalsProvider',            'NotifierProvider',        'Map<String,int> — transaksi/dompet/tabungan/utangPiutang'],
        ['transaksiProvider',                'AsyncNotifierProvider',   'List<Transaksi>'],
        ['dompetProvider',                   'AsyncNotifierProvider',   'List<Dompet>'],
        ['monthlySummaryProvider',            'FutureProvider.family',   'autoDispose + select + bulanPageProvider'],
        ['categorySummaryProvider',           'FutureProvider.family',   'autoDispose + select + bulanPageProvider'],
        ['budgetListProvider',               'FutureProvider.family',   'autoDispose + select + bulanPageProvider'],
        ['paginatedTransaksiByMonthProvider','FutureProvider.family',   'autoDispose + select + bulanPageProvider'],
        ['paginatedTransaksiByDateProvider', 'FutureProvider.family',   'autoDispose + select + perTanggalPageProvider'],
        ['transaksiByDateProvider',          'FutureProvider.family',   'select'],
        ['transaksiByMonthProvider',          'FutureProvider.family',   'select'],
        ['utangPiutangListProvider',        'FutureProvider',          'select'],
        ['tabunganImpianListProvider',       'FutureProvider',          'select'],
        ['pengaturanProvider',               'NotifierProvider',        'Pengaturan object'],
    ],
    col_widths=[Inches(2.5), Inches(2.0), Inches(2.7)]
)

add_section_heading(doc, '12.2 Reaktif Pattern', 2)
add_bullet(doc, 'Update signal system: Map dengan domain keys (transaksi/dompet/tabungan/utangPiutang)')
add_bullet(doc, 'Semua provider watch via .select() untuk tracking per-key Map')
add_bullet(doc, 'autoDispose.family untuk semua month/year parameterized providers')
add_bullet(doc, 'Pagination: perTanggalPageProvider, bulananPageProvider')
add_bullet(doc, 'Signal broadcast saat CRUD operasi')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  SECTION 13: I18N & HOME WIDGET
# ═══════════════════════════════════════════════════════════════
add_section_heading(doc, '13. I18N, HOME WIDGET & ERROR HANDLING', 1)

add_section_heading(doc, '13.1 Localization (l10n/)', 2)
add_bullet(doc, 'Indonesian (id_ID) — default')
add_bullet(doc, 'English (en_US)')
add_bullet(doc, 'Flutter Localizations delegates')
add_bullet(doc, 'Locale resolution: default ke Indonesian')

add_section_heading(doc, '13.2 Android Home Widget (home_widget_service.dart)', 2)
add_bullet(doc, 'Update Android home screen widget dengan saldo bulanan')
add_bullet(doc, 'Non-blocking (fail silently)')
add_bullet(doc, 'Dipanggil via ref.listen pada transaksiProvider')

add_section_heading(doc, '13.3 Error Handling (error_service.dart)', 2)
add_bullet(doc, 'Centralized error recording service')
add_bullet(doc, 'Silent catch di semua service layer')
add_bullet(doc, 'Non-blocking notifications / fallback behavior')

# ─── Save ────────────────────────────────────────────────────
out_dir = os.path.join(os.path.dirname(__file__))
out_path = os.path.join(out_dir, 'DompetKu_Fitur_Report.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
